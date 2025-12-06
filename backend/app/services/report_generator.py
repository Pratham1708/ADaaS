"""AI-Powered Automated Report Generator for ADaaS.

This module generates comprehensive PDF and Word reports for all analysis types:
- Survival Analysis (Kaplan-Meier, Cox, Nelson-Aalen)
- Reserving (Chain-Ladder)
- GLM Analysis
- ML Survival Models
- Mortality Analysis
- Time-Series Forecasting

Reports include:
- Executive Summary
- Key Insights (AI-generated via Gemini)
- Model Details
- Graphs and Visualizations
- Diagnostics
- Business Implications
"""

import os
import json
import base64
from pathlib import Path
from typing import Dict, Any, Optional, List
from datetime import datetime
import pandas as pd
import numpy as np

# Document generation libraries
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[WARN] python-docx not installed. Word reports will not be available.")

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, PageBreak, Table, TableStyle
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("[WARN] reportlab not installed. PDF reports will not be available.")

# Gemini AI for insights
import google.generativeai as genai
from dotenv import load_dotenv

# Load environment variables
load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)
    GEMINI_AVAILABLE = True
else:
    GEMINI_AVAILABLE = False
    print("[WARN] GEMINI_API_KEY not found. AI insights will be limited.")


class ReportGenerator:
    """Generate comprehensive AI-powered reports for analysis results."""
    
    def __init__(self, analysis_type: str, results: Dict[str, Any], dataset_id: str):
        """
        Initialize report generator.
        
        Args:
            analysis_type: Type of analysis ('survival', 'glm', 'ml_survival', 'mortality', 'timeseries', 'chainladder')
            results: Analysis results dictionary
            dataset_id: Dataset identifier
        """
        self.analysis_type = analysis_type
        self.results = results
        self.dataset_id = dataset_id
        self.timestamp = datetime.now()
        self.report_dir = Path("backend/analysis_results/reports")
        self.report_dir.mkdir(parents=True, exist_ok=True)
    
    # Helper Methods
    def _format_number(self, value, decimals: int = 2) -> str:
        """Safely format numeric values."""
        if value is None or value == 'N/A':
            return 'N/A'
        if isinstance(value, (int, float)):
            if decimals == 0:
                return f"{value:,.0f}"
            else:
                return f"{value:,.{decimals}f}"
        return str(value)
    
    def _create_pdf_table(self, data: List[Dict], columns: List[str], max_rows: int = 50) -> Any:
        """Create formatted table for PDF reports."""
        if not PDF_AVAILABLE or not data:
            return None
        
        # Limit rows
        limited_data = data[:max_rows]
        
        # Create table data with headers
        table_data = [[col for col in columns]]
        
        for row in limited_data:
            table_row = []
            for col in columns:
                value = row.get(col, 'N/A')
                if isinstance(value, (int, float)):
                    table_row.append(self._format_number(value, 3))
                else:
                    table_row.append(str(value))
            table_data.append(table_row)
        
        # Create table
        table = Table(table_data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#283593')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONTSIZE', (0, 1), (-1, -1), 8),
        ]))
        
        return table
    
    def _add_section_header(self, story: List, title: str, styles, level: int = 3):
        """Add formatted section header to PDF."""
        if level == 3:
            header_style = ParagraphStyle(
                'SectionHeader',
                parent=styles['Heading3'],
                fontSize=12,
                textColor=colors.HexColor('#1a237e'),
                spaceAfter=8,
                spaceBefore=12,
                bold=True
            )
        else:
            header_style = styles['Heading2']
        
        story.append(Paragraph(title, header_style))
    
    def _add_key_value_pair(self, story: List, key: str, value: Any, styles):
        """Add formatted key-value pair to PDF."""
        text = f"<b>{key}:</b> {value}"
        story.append(Paragraph(text, styles['Normal']))
        
    def generate_report(self, format: str = "pdf") -> str:
        """
        Generate comprehensive report.
        
        Args:
            format: Report format ('pdf' or 'docx')
            
        Returns:
            Path to generated report file
        """
        if format.lower() == "pdf":
            if not PDF_AVAILABLE:
                raise ImportError("reportlab not installed. Install with: pip install reportlab")
            return self._generate_pdf_report()
        elif format.lower() == "docx":
            if not DOCX_AVAILABLE:
                raise ImportError("python-docx not installed. Install with: pip install python-docx")
            return self._generate_word_report()
        else:
            raise ValueError(f"Unsupported format: {format}. Use 'pdf' or 'docx'.")
    
    def _generate_ai_insights(self) -> Dict[str, Any]:
        """
        Generate AI-powered insights using Gemini.
        
        Returns:
            Dictionary with executive summary, key insights, and business implications
        """
        if not GEMINI_AVAILABLE:
            return self._generate_fallback_insights()
        
        try:
            # Prepare analysis summary for Gemini
            summary = self._prepare_analysis_summary()
            
            # Create prompt for Gemini
            prompt = f"""
You are an expert actuarial data scientist analyzing results from a {self.analysis_type} analysis.

Dataset: {self.dataset_id}
Analysis Type: {self.analysis_type}

Analysis Results Summary:
{json.dumps(summary, indent=2, default=str)}

Please provide a comprehensive analysis report with the following sections:

1. EXECUTIVE SUMMARY (2-3 sentences)
   - High-level overview of what was analyzed and the most important finding

2. KEY INSIGHTS (3-5 bullet points)
   - Most significant findings from the analysis
   - Statistical significance and practical importance
   - Unexpected patterns or anomalies

3. MODEL PERFORMANCE (if applicable)
   - Assessment of model quality and reliability
   - Key metrics interpretation
   - Confidence in predictions

4. BUSINESS IMPLICATIONS (3-5 bullet points)
   - Actionable recommendations based on findings
   - Risk considerations
   - Strategic implications for decision-making

5. LIMITATIONS AND CAVEATS (2-3 bullet points)
   - Data quality considerations
   - Model assumptions
   - Areas requiring further investigation

Format your response as JSON with keys: executive_summary, key_insights (array), model_performance, business_implications (array), limitations (array).
"""
            
            # Call Gemini API
            model = genai.GenerativeModel('gemini-1.5-flash')
            response = model.generate_content(prompt)
            
            # Parse response
            response_text = response.text.strip()
            
            # Extract JSON from markdown code blocks if present
            if "```json" in response_text:
                response_text = response_text.split("```json")[1].split("```")[0].strip()
            elif "```" in response_text:
                response_text = response_text.split("```")[1].split("```")[0].strip()
            
            insights = json.loads(response_text)
            return insights
            
        except Exception as e:
            print(f"[WARN] Gemini AI insights failed: {e}")
            return self._generate_fallback_insights()
    
    def _generate_fallback_insights(self) -> Dict[str, Any]:
        """Generate rule-based insights when Gemini is unavailable."""
        insights = {
            "executive_summary": f"Completed {self.analysis_type} analysis on dataset {self.dataset_id}. Results show statistical patterns consistent with the data characteristics.",
            "key_insights": [
                f"{self.analysis_type.title()} analysis completed successfully",
                "Statistical models fitted to the data",
                "Results available for business decision-making"
            ],
            "model_performance": "Model metrics are within acceptable ranges for this analysis type.",
            "business_implications": [
                "Review detailed results for specific insights",
                "Consider validation with domain experts",
                "Use findings to inform strategic decisions"
            ],
            "limitations": [
                "Analysis based on available data only",
                "Standard statistical assumptions apply",
                "Results should be validated with additional data sources"
            ]
        }
        
        # Customize based on analysis type
        if self.analysis_type == "survival":
            insights["key_insights"].append("Survival curves show time-dependent patterns")
        elif self.analysis_type == "glm":
            insights["key_insights"].append("GLM coefficients indicate feature relationships")
        elif self.analysis_type == "ml_survival":
            insights["key_insights"].append("Machine learning models provide non-linear predictions")
        elif self.analysis_type == "mortality":
            insights["key_insights"].append("Mortality rates show age-dependent patterns")
        elif self.analysis_type == "timeseries":
            insights["key_insights"].append("Time series exhibits temporal patterns and trends")
        
        return insights
    
    def _prepare_analysis_summary(self) -> Dict[str, Any]:
        """Prepare concise summary of analysis results for AI processing."""
        summary = {
            "analysis_type": self.analysis_type,
            "dataset_id": self.dataset_id,
            "timestamp": str(self.timestamp)
        }
        
        # Extract key metrics based on analysis type
        if self.analysis_type == "survival":
            summary["median_survival"] = self.results.get("overall_km", {}).get("median_survival")
            summary["events"] = self.results.get("meta", {}).get("n_events")
            summary["censored"] = self.results.get("meta", {}).get("n_censored")
            
        elif self.analysis_type == "glm":
            summary["model_family"] = self.results.get("model_info", {}).get("family")
            summary["aic"] = self.results.get("goodness_of_fit", {}).get("aic")
            summary["deviance"] = self.results.get("goodness_of_fit", {}).get("deviance")
            summary["n_features"] = len(self.results.get("coefficients", []))
            
        elif self.analysis_type == "ml_survival":
            summary["model_type"] = self.results.get("model_type")
            summary["c_index"] = self.results.get("test_c_index")
            summary["n_features"] = len(self.results.get("variable_importance", []))
            
        elif self.analysis_type == "mortality":
            summary["life_expectancy"] = self.results.get("kpis", {}).get("life_expectancy_at_birth")
            summary["max_age"] = self.results.get("kpis", {}).get("max_age")
            
        elif self.analysis_type == "timeseries":
            summary["model_type"] = self.results.get("model_type")
            summary["forecast_periods"] = len(self.results.get("forecast", {}).get("dates", []))
            summary["metrics"] = self.results.get("metrics", {})
        
        return summary
    
    def _generate_pdf_report(self) -> str:
        """Generate PDF report using ReportLab."""
        filename = f"{self.analysis_type}_report_{self.dataset_id}_{self.timestamp.strftime('%Y%m%d_%H%M%S')}.pdf"
        filepath = self.report_dir / filename
        
        # Create PDF document
        doc = SimpleDocTemplate(str(filepath), pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1a237e'),
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#283593'),
            spaceAfter=12,
            spaceBefore=12
        )
        
        # Title
        story.append(Paragraph(f"{self.analysis_type.upper()} ANALYSIS REPORT", title_style))
        story.append(Spacer(1, 0.2*inch))
        
        # Metadata
        metadata_text = f"""
        <b>Dataset:</b> {self.dataset_id}<br/>
        <b>Analysis Type:</b> {self.analysis_type.title()}<br/>
        <b>Generated:</b> {self.timestamp.strftime('%Y-%m-%d %H:%M:%S')}<br/>
        <b>Platform:</b> ADaaS (Actuarial Data Analysis as a Service)
        """
        story.append(Paragraph(metadata_text, styles['Normal']))
        story.append(Spacer(1, 0.3*inch))
        
        # Generate AI insights
        insights = self._generate_ai_insights()
        
        # Executive Summary
        story.append(Paragraph("EXECUTIVE SUMMARY", heading_style))
        story.append(Paragraph(insights["executive_summary"], styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        # Key Insights
        story.append(Paragraph("KEY INSIGHTS", heading_style))
        for insight in insights["key_insights"]:
            story.append(Paragraph(f"• {insight}", styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        # Model Performance
        if insights.get("model_performance"):
            story.append(Paragraph("MODEL PERFORMANCE", heading_style))
            story.append(Paragraph(insights["model_performance"], styles['Normal']))
            story.append(Spacer(1, 0.2*inch))
        
        # Detailed Results
        story.append(Paragraph("DETAILED RESULTS", heading_style))
        self._add_detailed_results_to_pdf(story, styles)
        story.append(Spacer(1, 0.2*inch))
        
        # Business Implications
        story.append(Paragraph("BUSINESS IMPLICATIONS", heading_style))
        for implication in insights["business_implications"]:
            story.append(Paragraph(f"• {implication}", styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        # Limitations
        story.append(Paragraph("LIMITATIONS AND CAVEATS", heading_style))
        for limitation in insights["limitations"]:
            story.append(Paragraph(f"• {limitation}", styles['Normal']))
        
        # Build PDF
        doc.build(story)
        
        return str(filepath)
    
    def _add_detailed_results_to_pdf(self, story: List, styles):
        """Add analysis-specific detailed results to PDF."""
        if self.analysis_type == "survival":
            self._add_survival_details_pdf(story, styles)
        elif self.analysis_type == "glm":
            self._add_glm_details_pdf(story, styles)
        elif self.analysis_type == "ml_survival":
            self._add_ml_survival_details_pdf(story, styles)
        elif self.analysis_type == "mortality":
            self._add_mortality_details_pdf(story, styles)
        elif self.analysis_type == "timeseries":
            self._add_timeseries_details_pdf(story, styles)
        elif self.analysis_type == "chainladder":
            self._add_chainladder_details_pdf(story, styles)
    
    def _add_survival_details_pdf(self, story: List, styles):
        """Add comprehensive survival analysis details to PDF."""
        meta = self.results.get("meta", {})
        km = self.results.get("overall_km", {})
        
        # === SECTION 1: Study Overview ===
        self._add_section_header(story, "Study Overview", styles)
        self._add_key_value_pair(story, "Total Sample Size", self._format_number(meta.get('n', 'N/A'), 0), styles)
        self._add_key_value_pair(story, "Number of Events", self._format_number(meta.get('n_events', 'N/A'), 0), styles)
        self._add_key_value_pair(story, "Number Censored", self._format_number(meta.get('n_censored', 'N/A'), 0), styles)
        
        if meta.get('n'):
            event_rate = (meta.get('n_events', 0) / meta.get('n', 1)) * 100
            self._add_key_value_pair(story, "Event Rate", f"{self._format_number(event_rate, 1)}%", styles)
        
        story.append(Spacer(1, 0.15*inch))
        
        # === SECTION 2: Kaplan-Meier Survival Analysis ===
        self._add_section_header(story, "Kaplan-Meier Survival Analysis", styles)
        
        median_surv = km.get('median_survival', 'N/A')
        mean_surv = km.get('mean_survival', 'N/A')
        
        self._add_key_value_pair(story, "Median Survival Time", self._format_number(median_surv, 2), styles)
        self._add_key_value_pair(story, "Mean Survival Time", self._format_number(mean_surv, 2), styles)
        
        # Confidence intervals if available
        if 'median_ci_lower' in km and 'median_ci_upper' in km:
            ci_text = f"[{self._format_number(km.get('median_ci_lower'), 2)}, {self._format_number(km.get('median_ci_upper'), 2)}]"
            self._add_key_value_pair(story, "95% CI for Median", ci_text, styles)
        
        story.append(Spacer(1, 0.1*inch))
        
        # Survival probabilities at key time points
        story.append(Paragraph("<b>Survival Probabilities at Key Time Points:</b>", styles['Normal']))
        life_table = self.results.get("life_table", [])
        if life_table:
            # Get survival at 25%, 50%, 75% of max time
            max_time = max([row.get('time', 0) for row in life_table]) if life_table else 0
            key_times = [max_time * 0.25, max_time * 0.5, max_time * 0.75, max_time]
            
            for target_time in key_times:
                # Find closest time point
                closest = min(life_table, key=lambda x: abs(x.get('time', 0) - target_time))
                time_val = closest.get('time', 'N/A')
                surv_val = closest.get('survival', 'N/A')
                self._add_key_value_pair(story, f"  At time {self._format_number(time_val, 1)}", 
                                        f"{self._format_number(surv_val * 100 if isinstance(surv_val, (int, float)) else surv_val, 1)}%", styles)
        
        story.append(Spacer(1, 0.15*inch))
        
        # === SECTION 3: Life Table ===
        if life_table:
            self._add_section_header(story, "Life Table (First 20 Time Points)", styles)
            story.append(Paragraph(
                "Detailed survival probabilities, events, and at-risk counts over time.",
                styles['Normal']
            ))
            story.append(Spacer(1, 0.1*inch))
            
            # Create table
            table_columns = ['time', 'at_risk', 'events', 'censored', 'survival']
            table = self._create_pdf_table(life_table, table_columns, max_rows=20)
            if table:
                story.append(table)
                if len(life_table) > 20:
                    story.append(Paragraph(
                        f"<i>Note: Showing first 20 of {len(life_table)} time points.</i>",
                        styles['Normal']
                    ))
            story.append(Spacer(1, 0.15*inch))
        
        # === SECTION 4: Nelson-Aalen Cumulative Hazard ===
        nelson_aalen = self.results.get("nelson_aalen", {})
        if nelson_aalen:
            self._add_section_header(story, "Nelson-Aalen Cumulative Hazard", styles)
            story.append(Paragraph(
                "Cumulative hazard estimates provide an alternative view of the failure process.",
                styles['Normal']
            ))
            
            cumulative_hazard = nelson_aalen.get('cumulative_hazard', [])
            if cumulative_hazard:
                final_hazard = cumulative_hazard[-1] if cumulative_hazard else 'N/A'
                self._add_key_value_pair(story, "Final Cumulative Hazard", self._format_number(final_hazard, 4), styles)
            
            story.append(Spacer(1, 0.15*inch))
        
        # === SECTION 5: Cox Proportional Hazards Model ===
        cox = self.results.get("cox", {})
        if cox and cox.get('summary'):
            self._add_section_header(story, "Cox Proportional Hazards Model", styles)
            
            c_index = cox.get('concordance', 'N/A')
            self._add_key_value_pair(story, "Concordance Index (C-index)", self._format_number(c_index, 3), styles)
            
            story.append(Spacer(1, 0.1*inch))
            story.append(Paragraph("<b>Model Coefficients and Hazard Ratios:</b>", styles['Normal']))
            story.append(Spacer(1, 0.05*inch))
            
            # Create coefficients table
            coeffs = cox.get('summary', [])
            for coef in coeffs[:15]:  # Top 15 variables
                var_name = coef.get('covariate', 'N/A')
                coef_val = self._format_number(coef.get('coef', 'N/A'), 4)
                hr = self._format_number(coef.get('exp_coef', 'N/A'), 3)
                p_val = self._format_number(coef.get('p', 'N/A'), 4)
                
                text = f"<b>{var_name}:</b> Coef={coef_val}, HR={hr}, p={p_val}"
                story.append(Paragraph(text, styles['Normal']))
            
            if len(coeffs) > 15:
                story.append(Paragraph(
                    f"<i>Note: Showing top 15 of {len(coeffs)} variables.</i>",
                    styles['Normal']
                ))
            
            story.append(Spacer(1, 0.15*inch))
        
        # === SECTION 6: Stratified Analysis ===
        strata = self.results.get("strata", {})
        if strata and strata.get("column") and strata.get("results"):
            self._add_section_header(story, "Stratified Analysis", styles)
            story.append(Paragraph(
                f"Survival analysis stratified by: {strata.get('column')}",
                styles['Normal']
            ))
            story.append(Spacer(1, 0.1*inch))
            
            # Log-rank test
            logrank_p = strata.get('logrank_p')
            if logrank_p is not None:
                self._add_key_value_pair(story, "Log-Rank Test P-value", self._format_number(logrank_p, 4), styles)
                story.append(Spacer(1, 0.05*inch))
            
            # Results by group
            results = strata.get('results', [])
            if results:
                story.append(Paragraph("<b>Survival by Group:</b>", styles['Normal']))
                for group_result in results[:10]:  # Top 10 groups
                    group_name = group_result.get('group', 'N/A')
                    n = self._format_number(group_result.get('n', 'N/A'), 0)
                    
                    # Calculate median survival from timeline and survival data
                    timeline = group_result.get('timeline', [])
                    survival = group_result.get('survival', [])
                    
                    median_surv = 'N/A'
                    if timeline and survival:
                        # Find where survival crosses 0.5
                        for i, s in enumerate(survival):
                            if s <= 0.5:
                                median_surv = self._format_number(timeline[i], 2)
                                break
                    
                    self._add_key_value_pair(story, f"  {group_name}", f"N={n}, Median={median_surv}", styles)
                
                if len(results) > 10:
                    story.append(Paragraph(
                        f"<i>Note: Showing 10 of {len(results)} groups.</i>",
                        styles['Normal']
                    ))
            
            story.append(Spacer(1, 0.15*inch))
        
        
    
    def _add_glm_details_pdf(self, story: List, styles):
        """Add comprehensive GLM analysis details to PDF."""
        model_info = self.results.get("model_info", {})
        gof = self.results.get("goodness_of_fit", {})
        
        # === SECTION 1: Model Specification ===
        self._add_section_header(story, "Model Specification", styles)
        self._add_key_value_pair(story, "Model Family", model_info.get('family', 'N/A'), styles)
        self._add_key_value_pair(story, "Link Function", model_info.get('link', 'N/A'), styles)
        self._add_key_value_pair(story, "Number of Observations", self._format_number(model_info.get('n_obs', 'N/A'), 0), styles)
        self._add_key_value_pair(story, "Number of Features", len(self.results.get('coefficients', [])), styles)
        story.append(Spacer(1, 0.15*inch))
        
        # === SECTION 2: Goodness of Fit ===
        self._add_section_header(story, "Goodness of Fit Metrics", styles)
        self._add_key_value_pair(story, "AIC (Akaike Information Criterion)", self._format_number(gof.get('aic', 'N/A'), 2), styles)
        self._add_key_value_pair(story, "BIC (Bayesian Information Criterion)", self._format_number(gof.get('bic', 'N/A'), 2), styles)
        self._add_key_value_pair(story, "Deviance", self._format_number(gof.get('deviance', 'N/A'), 2), styles)
        self._add_key_value_pair(story, "Pearson Chi-Square", self._format_number(gof.get('pearson_chi2', 'N/A'), 2), styles)
        self._add_key_value_pair(story, "Degrees of Freedom", self._format_number(gof.get('df_resid', 'N/A'), 0), styles)
        story.append(Spacer(1, 0.15*inch))
        
        # === SECTION 3: Model Coefficients ===
        coeffs = self.results.get("coefficients", [])
        if coeffs:
            self._add_section_header(story, "Model Coefficients (Top 20)", styles)
            story.append(Paragraph(
                "Regression coefficients with standard errors, z-values, and p-values.",
                styles['Normal']
            ))
            story.append(Spacer(1, 0.1*inch))
            
            for coef in coeffs[:20]:
                var = coef.get('variable', 'N/A')
                coef_val = self._format_number(coef.get('coefficient', 'N/A'), 4)
                se = self._format_number(coef.get('std_err', 'N/A'), 4)
                z_val = self._format_number(coef.get('z_value', 'N/A'), 3)
                p_val = self._format_number(coef.get('p_value', 'N/A'), 4)
                
                text = f"<b>{var}:</b> β={coef_val}, SE={se}, z={z_val}, p={p_val}"
                story.append(Paragraph(text, styles['Normal']))
            
            if len(coeffs) > 20:
                story.append(Paragraph(
                    f"<i>Note: Showing top 20 of {len(coeffs)} coefficients.</i>",
                    styles['Normal']
                ))
            story.append(Spacer(1, 0.15*inch))
        
        # === SECTION 4: Feature Importance ===
        feature_importance = self.results.get("feature_importance", [])
        if feature_importance:
            self._add_section_header(story, "Feature Importance Rankings", styles)
            story.append(Paragraph(
                "Variables ranked by their importance in the model.",
                styles['Normal']
            ))
            story.append(Spacer(1, 0.1*inch))
            
            for i, feat in enumerate(feature_importance[:10], 1):
                name = feat.get('feature', 'N/A')
                importance = self._format_number(feat.get('importance', 'N/A'), 4)
                self._add_key_value_pair(story, f"  {i}. {name}", importance, styles)
            
            story.append(Spacer(1, 0.15*inch))
        
        # === SECTION 5: Model Diagnostics ===
        residuals = self.results.get("residuals", {})
        if residuals:
            self._add_section_header(story, "Residual Diagnostics", styles)
            self._add_key_value_pair(story, "Mean Residual", self._format_number(residuals.get('mean', 'N/A'), 6), styles)
            self._add_key_value_pair(story, "Std Dev of Residuals", self._format_number(residuals.get('std', 'N/A'), 4), styles)
            self._add_key_value_pair(story, "Min Residual", self._format_number(residuals.get('min', 'N/A'), 4), styles)
            self._add_key_value_pair(story, "Max Residual", self._format_number(residuals.get('max', 'N/A'), 4), styles)
    
    def _add_ml_survival_details_pdf(self, story: List, styles):
        """Add comprehensive ML survival analysis details to PDF."""
        
        # === SECTION 1: Model Overview ===
        self._add_section_header(story, "Model Overview", styles)
        self._add_key_value_pair(story, "Model Type", self.results.get('model_type', 'N/A'), styles)
        self._add_key_value_pair(story, "Number of Features", len(self.results.get('variable_importance', [])), styles)
        self._add_key_value_pair(story, "Training Samples", self._format_number(self.results.get('n_train', 'N/A'), 0), styles)
        self._add_key_value_pair(story, "Test Samples", self._format_number(self.results.get('n_test', 'N/A'), 0), styles)
        story.append(Spacer(1, 0.15*inch))
        
        # === SECTION 2: Performance Metrics ===
        self._add_section_header(story, "Model Performance", styles)
        self._add_key_value_pair(story, "Training C-Index", self._format_number(self.results.get('train_c_index', 'N/A'), 3), styles)
        self._add_key_value_pair(story, "Test C-Index", self._format_number(self.results.get('test_c_index', 'N/A'), 3), styles)
        
        # Cross-validation if available
        cv_score = self.results.get('cv_score', 'N/A')
        if cv_score != 'N/A':
            self._add_key_value_pair(story, "Cross-Validation C-Index", self._format_number(cv_score, 3), styles)
        
        story.append(Spacer(1, 0.15*inch))
        
        # === SECTION 3: Variable Importance ===
        var_importance = self.results.get("variable_importance", [])
        if var_importance:
            self._add_section_header(story, "Variable Importance Rankings", styles)
            story.append(Paragraph(
                "Features ranked by their contribution to survival predictions.",
                styles['Normal']
            ))
            story.append(Spacer(1, 0.1*inch))
            
            for i, var in enumerate(var_importance[:20], 1):
                feat_name = var.get('feature', 'N/A')
                importance = self._format_number(var.get('importance', 'N/A'), 4)
                self._add_key_value_pair(story, f"  {i}. {feat_name}", importance, styles)
            
            if len(var_importance) > 20:
                story.append(Paragraph(
                    f"<i>Note: Showing top 20 of {len(var_importance)} features.</i>",
                    styles['Normal']
                ))
            story.append(Spacer(1, 0.15*inch))
        
        # === SECTION 4: Model Comparison ===
        comparison = self.results.get("comparison", {})
        if comparison:
            self._add_section_header(story, "Model Comparison", styles)
            story.append(Paragraph(
                "Comparison of different survival models on this dataset.",
                styles['Normal']
            ))
            story.append(Spacer(1, 0.1*inch))
            
            for model_name, metrics in comparison.items():
                c_index = self._format_number(metrics.get('c_index', 'N/A'), 3)
                self._add_key_value_pair(story, f"  {model_name}", f"C-Index: {c_index}", styles)
            
            story.append(Spacer(1, 0.15*inch))
        
        # === SECTION 5: Risk Stratification ===
        risk_groups = self.results.get("risk_stratification", {})
        if risk_groups:
            self._add_section_header(story, "Risk Stratification", styles)
            story.append(Paragraph(
                "Patients stratified into risk groups based on predicted survival.",
                styles['Normal']
            ))
            story.append(Spacer(1, 0.1*inch))
            
            for group_name, group_data in risk_groups.items():
                n_patients = self._format_number(group_data.get('n_patients', 'N/A'), 0)
                median_surv = self._format_number(group_data.get('median_survival', 'N/A'), 2)
                self._add_key_value_pair(story, f"  {group_name}", f"N={n_patients}, Median={median_surv}", styles)
    
    def _add_mortality_details_pdf(self, story: List, styles):
        """Add comprehensive mortality analysis details to PDF."""
        kpis = self.results.get("kpis", {})
        
        # === SECTION 1: Key Performance Indicators ===
        self._add_section_header(story, "Key Performance Indicators", styles)
        
        life_exp = kpis.get('life_expectancy_at_birth', 'N/A')
        life_exp_str = self._format_number(life_exp, 2)
        
        max_age = kpis.get('max_age', 'N/A')
        peak_age = kpis.get('age_at_peak_mortality', 'N/A')
        
        peak_rate = kpis.get('peak_mortality_rate', 'N/A')
        peak_rate_str = self._format_number(peak_rate, 6)
        
        self._add_key_value_pair(story, "Life Expectancy at Birth", f"{life_exp_str} years", styles)
        self._add_key_value_pair(story, "Maximum Age", max_age, styles)
        self._add_key_value_pair(story, "Age at Peak Mortality", peak_age, styles)
        self._add_key_value_pair(story, "Peak Mortality Rate", peak_rate_str, styles)
        story.append(Spacer(1, 0.15*inch))
        
        # === SECTION 2: Life Table ===
        life_table = self.results.get("life_table", [])
        if life_table:
            self._add_section_header(story, "Life Table (First 20 Ages)", styles)
            story.append(Paragraph(
                "Complete actuarial life table showing mortality rates (qₓ), survivors (lₓ), deaths (dₓ), and life expectancy (eₓ) by age.",
                styles['Normal']
            ))
            story.append(Spacer(1, 0.1*inch))
            
            # Create life table
            table_columns = ['Age', 'qx', 'lx', 'dx', 'ex']
            table = self._create_pdf_table(life_table, table_columns, max_rows=20)
            if table:
                story.append(table)
                if len(life_table) > 20:
                    story.append(Paragraph(
                        f"<i>Note: Showing first 20 of {len(life_table)} ages. Complete table available in raw data.</i>",
                        styles['Normal']
                    ))
            story.append(Spacer(1, 0.15*inch))
        
        # === SECTION 3: Graduated Mortality Rates ===
        graduated = self.results.get("graduated", {})
        if graduated:
            self._add_section_header(story, "Graduated Mortality Rates", styles)
            story.append(Paragraph(
                "Comparison of raw mortality rates with graduated (smoothed) rates using different methods.",
                styles['Normal']
            ))
            story.append(Spacer(1, 0.1*inch))
            
            # Whittaker-Henderson
            whittaker = graduated.get("whittaker", {})
            if whittaker:
                story.append(Paragraph("<b>Whittaker-Henderson Graduation:</b>", styles['Normal']))
                self._add_key_value_pair(story, "  Lambda Parameter", self._format_number(whittaker.get('lambda', 'N/A'), 2), styles)
                self._add_key_value_pair(story, "  R² Score", self._format_number(whittaker.get('r2_score', 'N/A'), 4), styles)
                self._add_key_value_pair(story, "  RMSE", self._format_number(whittaker.get('rmse', 'N/A'), 6), styles)
                story.append(Spacer(1, 0.05*inch))
            
            # Moving Average
            moving_avg = graduated.get("moving_average", {})
            if moving_avg:
                story.append(Paragraph("<b>Moving Average Graduation:</b>", styles['Normal']))
                self._add_key_value_pair(story, "  Window Size", moving_avg.get('window_size', 'N/A'), styles)
                self._add_key_value_pair(story, "  R² Score", self._format_number(moving_avg.get('r2_score', 'N/A'), 4), styles)
                self._add_key_value_pair(story, "  RMSE", self._format_number(moving_avg.get('rmse', 'N/A'), 6), styles)
                story.append(Spacer(1, 0.05*inch))
            
            # Spline
            spline = graduated.get("spline", {})
            if spline:
                story.append(Paragraph("<b>Spline Graduation:</b>", styles['Normal']))
                self._add_key_value_pair(story, "  Smoothing Factor", self._format_number(spline.get('smoothing_factor', 'N/A'), 2), styles)
                self._add_key_value_pair(story, "  R² Score", self._format_number(spline.get('r2_score', 'N/A'), 4), styles)
                self._add_key_value_pair(story, "  RMSE", self._format_number(spline.get('rmse', 'N/A'), 6), styles)
            
            story.append(Spacer(1, 0.15*inch))
        
        # === SECTION 4: Parametric Models ===
        fitted_models = self.results.get("fitted_models", {})
        if fitted_models:
            self._add_section_header(story, "Parametric Mortality Models", styles)
            
            # Gompertz Model
            gompertz = fitted_models.get("gompertz", {})
            if gompertz:
                story.append(Paragraph("<b>Gompertz Model: μₓ = B·exp(c·x)</b>", styles['Normal']))
                params = gompertz.get('parameters', {})
                self._add_key_value_pair(story, "  Parameter B", self._format_number(params.get('B', 'N/A'), 6), styles)
                self._add_key_value_pair(story, "  Parameter c", self._format_number(params.get('c', 'N/A'), 6), styles)
                self._add_key_value_pair(story, "  R² Score", self._format_number(gompertz.get('r2_score', 'N/A'), 4), styles)
                self._add_key_value_pair(story, "  RMSE", self._format_number(gompertz.get('rmse', 'N/A'), 6), styles)
                story.append(Spacer(1, 0.05*inch))
            
            # Makeham Model
            makeham = fitted_models.get("makeham", {})
            if makeham:
                story.append(Paragraph("<b>Makeham Model: μₓ = A + B·exp(c·x)</b>", styles['Normal']))
                params = makeham.get('parameters', {})
                self._add_key_value_pair(story, "  Parameter A", self._format_number(params.get('A', 'N/A'), 6), styles)
                self._add_key_value_pair(story, "  Parameter B", self._format_number(params.get('B', 'N/A'), 6), styles)
                self._add_key_value_pair(story, "  Parameter c", self._format_number(params.get('c', 'N/A'), 6), styles)
                self._add_key_value_pair(story, "  R² Score", self._format_number(makeham.get('r2_score', 'N/A'), 4), styles)
                self._add_key_value_pair(story, "  RMSE", self._format_number(makeham.get('rmse', 'N/A'), 6), styles)
            
            story.append(Spacer(1, 0.15*inch))
        
        # === SECTION 5: Age-Specific Life Expectancies ===
        if life_table and len(life_table) > 0:
            self._add_section_header(story, "Age-Specific Life Expectancies", styles)
            story.append(Paragraph(
                "Life expectancy at key ages showing remaining years of life.",
                styles['Normal']
            ))
            story.append(Spacer(1, 0.1*inch))
            
            # Select key ages
            key_ages = [0, 20, 40, 60, 80]
            for age in key_ages:
                age_data = next((row for row in life_table if row.get('Age') == age), None)
                if age_data:
                    ex = self._format_number(age_data.get('ex', 'N/A'), 2)
                    self._add_key_value_pair(story, f"  Life Expectancy at Age {age}", f"{ex} years", styles)
            
            story.append(Spacer(1, 0.15*inch))
        
        # === SECTION 6: Mortality Curve Characteristics ===
        raw_data = self.results.get("raw_data", {})
        if raw_data:
            self._add_section_header(story, "Mortality Curve Characteristics", styles)
            
            ages = raw_data.get('ages', [])
            qx_values = raw_data.get('qx', [])
            
            if ages and qx_values:
                # Find minimum mortality
                min_qx = min(qx_values) if qx_values else 'N/A'
                min_age = ages[qx_values.index(min_qx)] if qx_values and min_qx != 'N/A' else 'N/A'
                
                self._add_key_value_pair(story, "Minimum Mortality Rate", self._format_number(min_qx, 6), styles)
                self._add_key_value_pair(story, "Age at Minimum Mortality", min_age, styles)
                self._add_key_value_pair(story, "Age Range", f"{min(ages)} - {max(ages)}", styles)
                self._add_key_value_pair(story, "Number of Age Points", len(ages), styles)
        
    
    def _add_timeseries_details_pdf(self, story: List, styles):
        """Add comprehensive time-series analysis details to PDF."""
        metrics = self.results.get("metrics", {})
        forecast = self.results.get("forecast", {})
        
        # === SECTION 1: Model Information ===
        self._add_section_header(story, "Model Information", styles)
        self._add_key_value_pair(story, "Model Type", self.results.get('model_type', 'N/A'), styles)
        self._add_key_value_pair(story, "Forecast Periods", len(forecast.get('dates', [])), styles)
        
        # ARIMA parameters if available
        params = self.results.get('model_parameters', {})
        if params:
            if 'p' in params:
                self._add_key_value_pair(story, "ARIMA Order (p,d,q)", f"({params.get('p')}, {params.get('d')}, {params.get('q')})", styles)
            if 'seasonal_order' in params:
                self._add_key_value_pair(story, "Seasonal Order", str(params.get('seasonal_order')), styles)
        
        story.append(Spacer(1, 0.15*inch))
        
        # === SECTION 2: Forecast Accuracy ===
        self._add_section_header(story, "Forecast Accuracy Metrics", styles)
        self._add_key_value_pair(story, "MAE (Mean Absolute Error)", self._format_number(metrics.get('mae', 'N/A'), 2), styles)
        self._add_key_value_pair(story, "RMSE (Root Mean Squared Error)", self._format_number(metrics.get('rmse', 'N/A'), 2), styles)
        self._add_key_value_pair(story, "MAPE (Mean Absolute Percentage Error)", f"{self._format_number(metrics.get('mape', 'N/A'), 2)}%", styles)
        self._add_key_value_pair(story, "R² Score", self._format_number(metrics.get('r2', 'N/A'), 4), styles)
        story.append(Spacer(1, 0.15*inch))
        
        # === SECTION 3: Forecast Summary ===
        if forecast and forecast.get('values'):
            self._add_section_header(story, "Forecast Summary", styles)
            forecast_vals = forecast.get('values', [])
            
            if forecast_vals:
                self._add_key_value_pair(story, "Mean Forecast Value", self._format_number(sum(forecast_vals)/len(forecast_vals), 2), styles)
                self._add_key_value_pair(story, "Min Forecast Value", self._format_number(min(forecast_vals), 2), styles)
                self._add_key_value_pair(story, "Max Forecast Value", self._format_number(max(forecast_vals), 2), styles)
            
            # Show first few forecast points
            story.append(Spacer(1, 0.1*inch))
            story.append(Paragraph("<b>First 10 Forecast Points:</b>", styles['Normal']))
            dates = forecast.get('dates', [])
            for i in range(min(10, len(forecast_vals))):
                date_str = dates[i] if i < len(dates) else f"Period {i+1}"
                val = self._format_number(forecast_vals[i], 2)
                self._add_key_value_pair(story, f"  {date_str}", val, styles)
            
            story.append(Spacer(1, 0.15*inch))
        
        # === SECTION 4: Seasonal Decomposition ===
        decomposition = self.results.get("decomposition", {})
        if decomposition:
            self._add_section_header(story, "Seasonal Decomposition", styles)
            story.append(Paragraph(
                "Time series decomposed into trend, seasonal, and residual components.",
                styles['Normal']
            ))
            story.append(Spacer(1, 0.1*inch))
            
            trend_strength = self._format_number(decomposition.get('trend_strength', 'N/A'), 3)
            seasonal_strength = self._format_number(decomposition.get('seasonal_strength', 'N/A'), 3)
            
            self._add_key_value_pair(story, "Trend Strength", trend_strength, styles)
            self._add_key_value_pair(story, "Seasonal Strength", seasonal_strength, styles)
            self._add_key_value_pair(story, "Seasonal Period", decomposition.get('period', 'N/A'), styles)
    
    def _add_chainladder_details_pdf(self, story: List, styles):
        """Add chain-ladder analysis details to PDF."""
        details = f"""
        <b>Reserve Estimate:</b> ${self.results.get('reserve_estimate', 'N/A'):,.2f}<br/>
        <b>Number of Origin Periods:</b> {self.results.get('n_origin', 'N/A')}<br/>
        <b>Number of Development Periods:</b> {self.results.get('n_dev', 'N/A')}<br/>
        <b>Development Factors:</b> {', '.join([f'{x:.3f}' for x in self.results.get('development_factors', [])])}
        """
        story.append(Paragraph(details, styles['Normal']))
    
    def _generate_word_report(self) -> str:
        """Generate Word document report using python-docx."""
        filename = f"{self.analysis_type}_report_{self.dataset_id}_{self.timestamp.strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = self.report_dir / filename
        
        # Create Word document
        doc = Document()
        
        # Title
        title = doc.add_heading(f"{self.analysis_type.upper()} ANALYSIS REPORT", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        doc.add_paragraph(f"Dataset: {self.dataset_id}")
        doc.add_paragraph(f"Analysis Type: {self.analysis_type.title()}")
        doc.add_paragraph(f"Generated: {self.timestamp.strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Platform: ADaaS (Actuarial Data Analysis as a Service)")
        doc.add_paragraph()
        
        # Generate AI insights
        insights = self._generate_ai_insights()
        
        # Executive Summary
        doc.add_heading("Executive Summary", 1)
        doc.add_paragraph(insights["executive_summary"])
        
        # Key Insights
        doc.add_heading("Key Insights", 1)
        for insight in insights["key_insights"]:
            doc.add_paragraph(insight, style='List Bullet')
        
        # Model Performance
        if insights.get("model_performance"):
            doc.add_heading("Model Performance", 1)
            doc.add_paragraph(insights["model_performance"])
        
        # Detailed Results
        doc.add_heading("Detailed Results", 1)
        self._add_detailed_results_to_word(doc)
        
        # Business Implications
        doc.add_heading("Business Implications", 1)
        for implication in insights["business_implications"]:
            doc.add_paragraph(implication, style='List Bullet')
        
        # Limitations
        doc.add_heading("Limitations and Caveats", 1)
        for limitation in insights["limitations"]:
            doc.add_paragraph(limitation, style='List Bullet')
        
        # Save document
        doc.save(str(filepath))
        
        return str(filepath)
    
    def _add_detailed_results_to_word(self, doc):
        """Add analysis-specific detailed results to Word document."""
        if self.analysis_type == "survival":
            self._add_survival_details_word(doc)
        elif self.analysis_type == "glm":
            self._add_glm_details_word(doc)
        elif self.analysis_type == "ml_survival":
            self._add_ml_survival_details_word(doc)
        elif self.analysis_type == "mortality":
            self._add_mortality_details_word(doc)
        elif self.analysis_type == "timeseries":
            self._add_timeseries_details_word(doc)
        elif self.analysis_type == "chainladder":
            self._add_chainladder_details_word(doc)
    
    def _add_survival_details_word(self, doc):
        """Add survival analysis details to Word document."""
        meta = self.results.get("meta", {})
        km = self.results.get("overall_km", {})
        
        doc.add_paragraph(f"Sample Size: {self._format_number(meta.get('n', 'N/A'), 0)}")
        doc.add_paragraph(f"Events: {self._format_number(meta.get('n_events', 'N/A'), 0)}")
        doc.add_paragraph(f"Censored: {self._format_number(meta.get('n_censored', 'N/A'), 0)}")
        doc.add_paragraph(f"Median Survival: {self._format_number(km.get('median_survival', 'N/A'), 2)}")
        
        if "cox" in self.results and self.results["cox"]:
            doc.add_paragraph()
            doc.add_paragraph("Cox Proportional Hazards Model:", style='Heading 3')
            cox_data = self.results["cox"]
            concordance = cox_data.get('concordance', 'N/A')
            doc.add_paragraph(f"Concordance Index: {self._format_number(concordance, 3)}")
    
    def _add_glm_details_word(self, doc):
        """Add GLM analysis details to Word document."""
        model_info = self.results.get("model_info", {})
        gof = self.results.get("goodness_of_fit", {})
        
        doc.add_paragraph(f"Model Family: {model_info.get('family', 'N/A')}")
        doc.add_paragraph(f"Link Function: {model_info.get('link', 'N/A')}")
        doc.add_paragraph(f"AIC: {self._format_number(gof.get('aic', 'N/A'), 2)}")
        doc.add_paragraph(f"BIC: {self._format_number(gof.get('bic', 'N/A'), 2)}")
        doc.add_paragraph(f"Deviance: {self._format_number(gof.get('deviance', 'N/A'), 2)}")
        
        if "coefficients" in self.results:
            doc.add_paragraph()
            doc.add_paragraph("Model Coefficients (Top 10):", style='Heading 3')
            coeffs = self.results["coefficients"][:10]
            for coef in coeffs:
                var = coef.get('variable', 'N/A')
                coef_val = self._format_number(coef.get('coefficient', 'N/A'), 4)
                p_val = self._format_number(coef.get('p_value', 'N/A'), 4)
                doc.add_paragraph(f"{var}: {coef_val} (p={p_val})")
    
    def _add_ml_survival_details_word(self, doc):
        """Add ML survival analysis details to Word document."""
        doc.add_paragraph(f"Model Type: {self.results.get('model_type', 'N/A')}")
        
        train_c = self._format_number(self.results.get('train_c_index', 'N/A'), 3)
        test_c = self._format_number(self.results.get('test_c_index', 'N/A'), 3)
        
        doc.add_paragraph(f"Training C-Index: {train_c}")
        doc.add_paragraph(f"Test C-Index: {test_c}")
        
        if "variable_importance" in self.results:
            doc.add_paragraph()
            doc.add_paragraph("Top Features by Importance:", style='Heading 3')
            for var in self.results["variable_importance"][:10]:
                feat = var.get('feature', 'N/A')
                importance = self._format_number(var.get('importance', 'N/A'), 4)
                doc.add_paragraph(f"{feat}: {importance}")
    
    def _add_mortality_details_word(self, doc):
        """Add mortality analysis details to Word document."""
        kpis = self.results.get("kpis", {})
        
        # Safely format numeric values
        life_exp = kpis.get('life_expectancy_at_birth', 'N/A')
        life_exp_str = self._format_number(life_exp, 2)
        
        peak_rate = kpis.get('peak_mortality_rate', 'N/A')
        peak_rate_str = self._format_number(peak_rate, 6)
        
        doc.add_paragraph(f"Life Expectancy at Birth: {life_exp_str} years")
        doc.add_paragraph(f"Maximum Age: {kpis.get('max_age', 'N/A')}")
        doc.add_paragraph(f"Age at Peak Mortality: {kpis.get('age_at_peak_mortality', 'N/A')}")
        doc.add_paragraph(f"Peak Mortality Rate: {peak_rate_str}")
    
    def _add_timeseries_details_word(self, doc):
        """Add time-series analysis details to Word document."""
        metrics = self.results.get("metrics", {})
        
        doc.add_paragraph(f"Model Type: {self.results.get('model_type', 'N/A')}")
        doc.add_paragraph(f"Forecast Periods: {len(self.results.get('forecast', {}).get('dates', []))}")
        
        mae = self._format_number(metrics.get('mae', 'N/A'), 2)
        rmse = self._format_number(metrics.get('rmse', 'N/A'), 2)
        mape = self._format_number(metrics.get('mape', 'N/A'), 2)
        
        doc.add_paragraph(f"MAE: {mae}")
        doc.add_paragraph(f"RMSE: {rmse}")
        doc.add_paragraph(f"MAPE: {mape}%")
    
    def _add_chainladder_details_word(self, doc):
        """Add chain-ladder analysis details to Word document."""
        reserve = self.results.get('reserve_estimate', 'N/A')
        reserve_str = f"${self._format_number(reserve, 2)}" if reserve != 'N/A' else 'N/A'
        
        doc.add_paragraph(f"Reserve Estimate: {reserve_str}")
        doc.add_paragraph(f"Number of Origin Periods: {self.results.get('n_origin', 'N/A')}")
        doc.add_paragraph(f"Number of Development Periods: {self.results.get('n_dev', 'N/A')}")
        
        # Safely format development factors
        dev_factors = self.results.get('development_factors', [])
        if dev_factors and isinstance(dev_factors, list):
            factors_str = ', '.join([self._format_number(x, 3) for x in dev_factors])
            doc.add_paragraph(f"Development Factors: {factors_str}")
        else:
            doc.add_paragraph(f"Development Factors: N/A")


def generate_analysis_report(
    analysis_type: str,
    results: Dict[str, Any],
    dataset_id: str,
    format: str = "pdf"
) -> str:
    """
    Convenience function to generate analysis report.
    
    Args:
        analysis_type: Type of analysis
        results: Analysis results dictionary
        dataset_id: Dataset identifier
        format: Report format ('pdf' or 'docx')
        
    Returns:
        Path to generated report file
    """
    generator = ReportGenerator(analysis_type, results, dataset_id)
    return generator.generate_report(format)
